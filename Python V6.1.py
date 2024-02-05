import pandas as pd
import openpyxl
import pyinputplus as pyip
from openpyxl.chart import LineChart, Reference, BarChart
import matplotlib.pyplot as plt
import os
import xlwings as xw
import docx
import random
import subprocess
from openpyxl.styles import Font, PatternFill
import datetime, time


# In[2]:


def summary():
    #taking the time in the request
    startTime = time.time()
    print("Processing...")
    wb = openpyxl.load_workbook('spotify_songs.xlsx')
    sheet = wb['Sheet1']

    #creating a new wordsheet

    doc = docx.Document()
    title1 = "Spotify records"

    doc.add_heading(title1, 0)

    #taking max of columns and rows

    numrows = sheet.max_row
    numcolumns = sheet.max_column

    #adding title 1 in the word

    doc.add_paragraph(f"Our database is contain:  {numcolumns} columns and {numrows} rows")

    #creating a list of the columns'name    
    columnsname = []

    for i in range(1, numcolumns + 1):
        columnsname.append(sheet.cell(row=1, column=i).value)

    doc.add_paragraph(f"These columns will show the following information: {columnsname}")

    #creating a new title 
    genrescolumn = sheet['J']

    genrefind = []

    for cell in genrescolumn[1:]:
        genre = cell.value
        if genre:
            genrefind.append(genre)

    genreslist = list(set(genrefind))
    
    title2 = "Genre"

    doc.add_heading(title2, 0)
    doc.add_paragraph(f"These are: ")

    #creating a list
    j=1
    for k in genreslist:
        doc.add_paragraph(f"{j}. {k}")
        j+=1

    artistcolumn = sheet['C']

    #Creating the option to track of the artist

    artistfind = []

    for cell in artistcolumn[1:]:
        artist = cell.value
        if artist:
            artistfind.append(artist)

    artistlist = list(set(artistfind))

    #save 20 ramdom artist 

    random20artists = random.sample(artistfind, min(20, len(artistfind)))

    #save the information on the wordsheet

    title3 = "Artists"
    doc.add_heading(title3, 0)

    totalartist = len(artistlist)
    doc.add_paragraph(f"The number of artists is {totalartist} and 20 randomly selected artist names are: ")
    i=1
    for j in random20artists:   
        doc.add_paragraph(f"{i}. {j}")
        i+=1
    #save the document 
    doc.save('summary.docx')
    endTime = time.time()
    elapsedTime = round(endTime - startTime,2)
    print(f"The process took {elapsedTime} seconds")
    print("The wordsheet is ready, it will open automatically")

    #adding the option to open file automatically
    executablepath = r'C:\\Program Files\\Microsoft Office\\root\\Office16\\WINWORD.EXE'
    Opendoc = subprocess.Popen([executablepath, 'summary.docx'])


# In[3]:


def findTop10Tracks():

    startTime = time.time()
    print("Processing...")
    wb = pd.read_excel("spotify_songs.xlsx")

# Sort the data by column "track_popularity"
    sortedData = wb.sort_values(
        by=["track_popularity"], axis=0, ascending=False)
    sortedData.to_excel("spotify_songs.xlsx", index=False)

# Store data
    ranks = []
    songs = []

# Open Excel File with sorted data
    openFile = openpyxl.load_workbook("spotify_songs.xlsx")
    sheet = openFile.active

    # Iterate through each row and append values to the lists
    for row in sheet.iter_rows(min_row=0, values_only=True):
        # We will check if the song is not repeated and we will limit iteration to 10
        if row[1] not in songs and len(songs) < 11:
            songs.append(row[1])  # Column B
            ranks.append(row[3])  # Column D

    # Create a new workbook and add a new sheet
    new_wb = openpyxl.Workbook()
    new_sheet = new_wb.active
    new_sheet.title = "Top 10 Tracks Report"

    # Write the values to the new sheet
    for b_value, d_value in zip(songs, ranks):
        new_sheet.append([b_value, d_value])

    # Create a bar chart
    chart = BarChart()
    chart.title = "Top 10 Tracks"
    chart.x_axis.title = "Track Name"
    chart.y_axis.title = "Popularity"

    # Select data for the chart
    popularity_reference = Reference(new_sheet, min_col=2, min_row=1,
                                     max_col=2, max_row=len(songs))
    song_reference = Reference(
        new_sheet, min_col=1, min_row=2, max_row=len(songs))

    # Add data to the chart
    chart.add_data(popularity_reference, titles_from_data=True)
    chart.set_categories(song_reference)

    # Add the chart to the worksheet
    new_sheet.add_chart(chart, "D2")

    # Save the new workbook to a file
    new_wb.save("Top10Tracks.xlsx")
    endTime = time.time()
    elapsedTime = round(endTime - startTime,2)
    print("Data exported to Top10Tracks.xlsx")
    print(f"The process took {elapsedTime} seconds")


# In[4]:


def findDanceability():
    print("Processing...")
    startTime = time.time()

    # Load data from the Excel file
    fileName = 'spotify_songs.xlsx'
    df = pd.read_excel(fileName)

    # Find top 10 songs based on danceability
    topDanceability = df.sort_values(by='danceability', ascending=False).head(10)

    # Prepare a DataFrame with selected columns
    resultOfSongs = topDanceability[['track_name', 'danceability']]
    resultOfSongs.columns = ['Tracks', 'Danceable']

    # Save the DataFrame to an Excel file
    outputFile = 'reportOfDanceability.xlsx'
    resultOfSongs.to_excel(outputFile, index=False)

    # Read the Excel file for further processing
    df2 = pd.read_excel(outputFile)

    # Extract the rightmost two digits of 'Danceable', convert to integer, and update the DataFrame
    df2['rightMostTwo'] = df2['Danceable'].astype(str).str[-2:]
    df2['rightMostTwo'] = df2['rightMostTwo'].astype(int)
    df2 = df2.drop(columns=['Danceable'])
    df2 = df2.rename(columns={'rightMostTwo': 'Danceable'})

    # Save the updated DataFrame to the Excel file
    df2.to_excel(outputFile, index=False)

    # Truncate long track names for better visualization
    df2['Tracks'] = df2['Tracks'].apply(lambda x: x[:6] + '...' if len(x) > 6 else x)

    # Create a bar plot
    fig, ax = plt.subplots(figsize=(16, 12))
    df2.plot.bar(x='Tracks', y='Danceable', rot=0, ax=ax)
    plt.xticks(rotation='vertical')

    # Save the plot as an image
    plotFile = 'thePlot.png'
    plt.savefig(plotFile)
    plt.close(fig)

    # Load the Excel file and add the plot as an image
    wb = openpyxl.load_workbook(outputFile)
    ws = wb.active
    img = openpyxl.drawing.image.Image(plotFile)
    imgCell = 'A18'
    ws.add_image(img, imgCell)

    # Customize the title cell
    titleCellRange = ws['A14:C14']
    ws.merge_cells('A14:C14')
    titleCell = ws['A14']
    titleCell.value = "Top 10 Danceable Songs"
    titleCell.font = Font(bold=True, color='023f85')
    titleCell.fill = PatternFill(start_color='00B2E6AC', end_color='00B2E6AC', fill_type='solid')

    # Save the updated Excel file
    wb.save(outputFile)

    # Remove the temporary plot file
    os.remove("thePlot.png")

    endTime = time.time()
    elapsedTime = round(endTime - startTime, 2)
    print(f"Data exported to {outputFile}")
    print(f"The process took {elapsedTime} seconds")


# In[5]:
def topSongsByDecade():
    print("Processing...")
    # Start measuring execution time
    startTime = time.time()

    # Input file path
    filePath = 'spotify_songs.xlsx'

    # Read data from Excel file into a DataFrame
    dataFrame = pd.read_excel(filePath)

    # Convert 'track_album_release_date' to datetime and create a 'decade' column
    dataFrame['track_album_release_date'] = pd.to_datetime(dataFrame['track_album_release_date'], errors='coerce')
    dataFrame['decade'] = (dataFrame['track_album_release_date'].dt.year // 10) * 10

    # Filter data for songs released between 1960 and 2020, drop duplicates based on 'track_name'
    filteredData = dataFrame[(dataFrame['track_album_release_date'].dt.year >= 1960) &
                             (dataFrame['track_album_release_date'].dt.year <= 2020)]
    filteredData = filteredData.drop_duplicates(subset=['track_name'])

    # Sort the data by 'track_popularity' in descending order
    dfSorted = filteredData.sort_values(by='track_popularity', ascending=False)

    # Group by 'decade' and select the top 10 songs from each decade
    topSongsByDecade = dfSorted.groupby(['decade']).head(10)

    # Output file name
    file_name = 'topSongsByDecade.xlsx'

    # Write to Excel file with separate worksheets for each decade
    with pd.ExcelWriter(file_name, engine='openpyxl') as writer:
        for decade, group in topSongsByDecade.groupby('decade'):
            group[['track_name', 'track_popularity']].to_excel(writer, index=False, sheet_name=f'{decade - 1}-{decade}')

    # Stop measuring execution time
    endTime = time.time()
    executionTime = endTime - startTime

    # Print information about the process
    print(f"Exported the top ten songs of each decade to {file_name}, each decade in a separate worksheet.")
    print(f"This process took {executionTime:.2f} seconds")


# In[6]:
def workoutSongs():
    print("Processing...")
    startTime = time.time()
    wb = openpyxl.load_workbook("spotify_songs.xlsx")
    sheet = wb['Sheet1']

    # Create a new workbook and add a sheet
    newfile = openpyxl.Workbook()
    newsheet = newfile.active

    title1 = sheet['B1'].value
    title2 = sheet['L1'].value
    title3 = sheet['M1'].value

    # Write title to the new sheet
    newsheet.cell(row=1, column=1, value=title1)
    newsheet.cell(row=1, column=2, value=title2)
    newsheet.cell(row=1, column=3, value=title3)

    maxrow = sheet.max_row

    for i in range(2, maxrow + 1): 
        valueL = sheet['L' + str(i)].value
        valueM = sheet['M' + str(i)].value
        if valueL > 0.88 and valueM > 0.95:
       
            newsheet.append([sheet['B' + str(i)].value, valueL, valueM])


    # Create a line chart
    chart = LineChart()
    chart.title = "More Energy and Danciability songs"
    chart.x_axis.title = "Song's Name"
    chart.y_axis.title = "Values"


    #reference datatype
    data = Reference(newsheet, min_col=2, min_row=1, max_col=3, max_row = newsheet.max_row)
    chart.add_data(data, titles_from_data=True)
    
    data = Reference(newsheet, min_col=1, min_row=2, max_row = newsheet.max_row)
    chart.set_categories(data)
 
    newsheet.add_chart(chart, "F1")
        

    # Save the new workbook to a file
    newfile.save("workingOutSong.xlsx")
    endTime = time.time()
    elapsedTime = round(endTime - startTime,2)
    print("Table and grafic done")
    print("Data exported to workingOutSong.xlsx")
    print(f"The process took {elapsedTime} seconds")


# In[11]:


def genre():
    
    print("Processing...")
    startTime = time.time()

    wb = openpyxl.load_workbook("spotify_songs.xlsx")
    sheet = wb['Sheet1']
    #creating the variable to find in the columm
    genre1 ="pop"
    genre2 ="rap"
    genre3 = "latin"
    genre4 = "r&b"
    genre5 = "edm"
    genre6 = "rock"
    #variable to save the count
    countgenre1 = 0
    countgenre2 = 0
    countgenre3 = 0
    countgenre4 = 0
    countgenre5 = 0
    countgenre6 = 0

    #check rows per rows to find our variables.

    for i in range(2, sheet.max_row + 1):
    
        genre = sheet["J"+str(i)].value
    
        if genre == genre1:
            countgenre1 += 1
        elif genre == genre2:
            countgenre2 += 1
        elif genre == genre3:
            countgenre3 += 1
        elif genre == genre4:
            countgenre4 += 1
        elif genre == genre5:
            countgenre5 += 1
        elif genre == genre6:
            countgenre6 += 1
        
    #create a new file        
    newfile = openpyxl.Workbook()
    newsheet = newfile.active

    #create the title
    newsheet.cell(row=1, column=1, value="Genre")
    newsheet.cell(row=1, column=2, value="Count")

    #adding values
    newsheet.append([genre1, countgenre1])
    newsheet.append([genre2, countgenre2])
    newsheet.append([genre3, countgenre3])
    newsheet.append([genre4, countgenre4])
    newsheet.append([genre5, countgenre5])
    newsheet.append([genre6, countgenre6])


    # Create a bar chart
    chart = BarChart()
    chart.title = "Genre Counts"
    chart.x_axis.title = "Genre"
    chart.y_axis.title = "Count"

    # Reference datatype for chart data
    data = Reference(newsheet, min_col=2, min_row=1, max_col=2, max_row= newsheet.max_row)
    categories = Reference(newsheet, min_col=1, min_row=2, max_row=newsheet.max_row)

    chart.add_data(data, titles_from_data=True)
    chart.set_categories(categories)

    # Add the chart to the worksheet
    newsheet.add_chart(chart, "E3")
    #save the data 
    newfile.save("genre.xlsx")
    endTime = time.time()
    elapsedTime = round(endTime - startTime,2)
    print("Data exported to genre.xlsx")
    print(f"The process took {elapsedTime} seconds")


# In[7]:


def top10Artists():
    print('Processing...')
    startTime = time.time()
    #Creating pandas dataframe form excel file
    data=pd.read_excel('spotify_songs.xlsx')

    #Create table for top artists, based on # of tracks. head(10) method saves only first 10 artists
    artistCounts = data['track_artist'].value_counts().head(10).rename_axis('artists').reset_index(name='counts')

    #save table to excel top10_artists
    artistCounts.to_excel('top10_artists.xlsx')

    #Bar chart to visually represent the dara
    fig=plt.figure()
    x=artistCounts['artists']
    y=artistCounts['counts']
    plt.xticks(rotation=90) #rotate names of artists for readability
    plt.title("Top 10 popular artists")
    plt.bar(x,y)
    plt.grid=False

    #open and adding the chart to the existing xslx file (with table from previous step) using xlwings library
    wb=xw.Book("top10_artists.xlsx")
    sht=wb.sheets[0]
    sht.name='Table and plot'
    sht.pictures.add(
    fig,
    name='Top 10 artists',
    update=True,
    left=sht.range('E3').left,
    top=sht.range('E3').top,
    height=300,
    width=400)

    #save the file
    wb.save()
    endTime = time.time()
    elapsedTime = round(endTime - startTime,2)
    print("Data exported to top10_artists.xlsx")
    print(f"All process took {elapsedTime} seconds") 


# In[12]:

print('Welcome to Spotify analysis tool. Here are the options of available reports:')
userChoice='test'
#Menu with available reports
while(userChoice!='Exit'):
    userChoice = pyip.inputMenu(["Summary (word file)",
                            "Top 10 most popular songs Report",
                            "Top 10 most danceable songs Report",
                            "Top 10 most popular songs of each decade Report",
                             "Top 10 artists Report", 
                             "Top 10 workout Songs Report",
                                "Top genres Report",
                                "Exit"], prompt="\nPlease select a report from the options below: \n", numbered=True)
    

    if userChoice == "Top 10 most popular songs Report":
        findTop10Tracks()
    elif userChoice == "Top 10 most danceable songs Report":
        findDanceability()
    elif userChoice == "Top 10 most popular songs of each decade Report":
        topSongsByDecade()
    elif userChoice=="Top 10 artists Report":
        top10Artists()
    elif userChoice=="Top 10 workout Songs Report":
        workoutSongs()
    elif userChoice=="Top genres Report":
        genre()
    elif userChoice=="Summary (word file)":
        summary()
    elif userChoice=="Exit":
        print('Goodbye! Thank you for using our tool!')

